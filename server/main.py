import sys
from flask import Flask, request, send_file
from docx import Document
from docx.shared import Pt
from flask_cors import CORS
from docx2pdf import convert
import pythoncom
import threading
import logging
import datetime
import os
from concurrent_log_handler import ConcurrentRotatingFileHandler
from waitress import serve

app = Flask(__name__)
CORS(app)

log_handler = ConcurrentRotatingFileHandler('app.log', maxBytes=10000, backupCount=3)
log_handler.setLevel(logging.INFO)
app.logger.addHandler(log_handler)
app.logger.setLevel(logging.INFO)

lock = threading.Lock() 

def check_and_save_enrollment_number(eno, file_path):
    with lock:
        if os.path.exists(file_path):
            with open(file_path, "r") as f:
                existing_numbers = f.read().splitlines()
                if eno in existing_numbers:
                    return False  # Duplicate found
        with open(file_path, "a") as f:
            f.write(eno + "\n")  
    return True  

def generate_and_convert(data, number_str):
    pythoncom.CoInitialize()  
    
    document = Document("input_r.docx")
    
    fee_paid_in_sem = data["in_semester"]
    serial_year = "2024-2025"
    serial_year_int = 2024
    name = data['name']
    eno = data["enrollment_number"]
    acad_year = data["academic_year"]
    p_acad_year = data["previous_academic_year"]
    branch = data["branch"]
    amount = data["amount"]
    gender = data["gender"]
    hostel = data["hosteler"]
    mis = "Mr." if gender == "male" else "Ms."
    heshe = "He" if gender == "male" else "She"
    feepaid = data["fee_paid"]
    passing_year_1 = data["passing_year_1"]
    passing_sem_1 = data["passing_sem_1"]
    passing_sem_2 = data["passing_sem_2"]
    percentile_1 = data["percentile_1"]
    spi_1 = data["spi_1"]
    spi_2 = data["spi_2"]
    attempts_1 = data["attempts_1"]
    attempts_2 = data["attempts_2"]
    result_pass = ["PASS"]
    result_fail = ["FAIL"]
    xoxx = data['xoxx']
    today = datetime.date.today()
    placeholders = {
        "SERI": f"VPMP/MYSY/1084/R{number_str}/{serial_year}",
        "MIS": mis,
        "NAMEOFSTUDENT": name,
        "ENO": str(eno),
        "YOFATEN": str(serial_year_int - 1),
        "ACADYEAR": str(acad_year),
        "BRANCH": branch,
        "AMOUNTOFMONEY": str(amount),
        "PRYEAR": p_acad_year,
        "HESHE": heshe,
        "HOSTEL": hostel,
        "FEEPAID": feepaid,
        "FINSEM": fee_paid_in_sem,
        "PY_O": passing_year_1,
        "PSAM_O": passing_sem_1,
        "PSAM_S": passing_sem_2,
        "PERTILE_O": str(percentile_1),
        "SPI_O": spi_1,
        "SPI_S": spi_2,
        "ATM_O": attempts_1,
        "ATM_S": attempts_2,
        "PASS": result_pass[0],
        "FAIL": result_fail[0],
        "CURRD": today.strftime("%d/%m/%Y"),
        "XOXX": xoxx
    }

    for para in document.paragraphs:
        for run in para.runs:
            for key, value in placeholders.items():
                if key.strip() in run.text.strip() or key in run.text:
                    run.text = run.text.replace(key, value)
                    run.bold = True

    for para in document.paragraphs:
        for run in para.runs:
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(14)

    file_name_docx = f'files/re/{number_str}.docx'
    file_name_pdf = f'files/re/{number_str}.pdf'

    with lock:  
        document.save(file_name_docx)
        convert(file_name_docx, file_name_pdf)

    with open("counter_re.txt", "w") as f:
        f.write(str(int(number_str) + 1))

    pythoncom.CoUninitialize()  

    return file_name_pdf

@app.route('/test', methods=['POST'])
def handle_request():
    data = request.get_json()
    with open("./counter_re.txt", "r") as f:
        number_str = int(f.read().strip())

    eno = data["enrollment_number"]
    enrollments_file = "enrollments_re.txt"
    
    if not check_and_save_enrollment_number(eno, enrollments_file):
        return "Duplicate enrollment number error", 400

    try:
        file_path = generate_and_convert(data, number_str)
        return send_file(file_path, as_attachment=False, download_name=f"document_{number_str}.pdf", mimetype='application/pdf')
    except Exception as e:
        app.logger.error(f"Error generating PDF: {e}")
        return "Error generating PDF", 500

@app.route('/generate-doc', methods=['POST'])
def generate_word_doc():
    data = request.get_json()
    number_str = 0  

    try:
        with open("./counter.txt", "r") as f:
            content = f.read().strip()
            if content.isdigit():  
                number_str = int(content)
    except Exception as e:
        app.logger.error(f"Error reading/parsing counter.txt: {e}")

    try:
        with open("counter.txt", "w") as f:
            f.write(str(number_str + 1))
    except Exception as e:
        app.logger.error(f"Error writing to counter.txt: {e}")

    pythoncom.CoInitialize()

    eno = data['eno']
    enrollments_file = "enrollments_fresh.txt"

    if not check_and_save_enrollment_number(eno, enrollments_file):
        return "Duplicate enrollment number error", 400

    document = Document("input.docx")
    name = data['name']
    year = data['year']
    branch = data['branch']
    method = data['method']
    fee = data['fee']
    is_admitted = data['is_admitted']
    gender = data['gender']
    mis = "Mr." if gender == "male" else "Ms."
    heshe = "He" if gender == "male" else "She"
    today = datetime.date.today()
    placeholders = {
        "SERI": str(f"VPMP/MYSY/1083/F{number_str}/2024-2025"),
        "MIS": mis,
        "NAMEOFSTUDENT": name,
        "ENO": str(eno),
        "YEARINPUT": str(year),
        "BRANCH": branch,
        "METHOD": method,
        "HESHE": heshe,
        "AMOUNTOFFEE": str(fee),
        "ADMIT": is_admitted,
        "CURRD": f"{today.strftime('%d/%m/%Y')}"
    }

    for para in document.paragraphs:
        for run in para.runs:
            for key, value in placeholders.items():
                if key.strip() in run.text.strip() or key in run.text:
                    run.text = run.text.replace(key, value)
                    run.bold = True

    for para in document.paragraphs:
        for run in para.runs:
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(14)

    try:
        document.save(f'./files/fresh/{number_str}.docx')
        convert(f'./files/fresh/{number_str}.docx', f'./files/fresh/{number_str}.pdf')
    except Exception as e:
        app.logger.error(f"Error saving/convert document: {e}")

    pythoncom.CoUninitialize()

    try:
        return send_file(f'./files/fresh/{number_str}.pdf', as_attachment=True)
    except Exception as e:
        app.logger.error(f"Error sending file: {e}")
        return str(e), 500  

@app.route("/shutdown", methods=['GET'])
def shutdown():
    func = request.environ.get('werkzeug.server.shutdown')
    if func is None:
        raise RuntimeError('Not running with the Werkzeug Server')
    func()
    return 'Server shutting down...'

if __name__ == '__main__':
    port = int(sys.argv[1]) if len(sys.argv) > 1 else 5000
    serve(app, host='0.0.0.0', port=port)
